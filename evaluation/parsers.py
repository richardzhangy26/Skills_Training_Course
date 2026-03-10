"""
文件解析器 - 支持docx/md/txt教师文档和json/txt对话日志
"""

import json
import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .types import (
    TeacherDocument,
    DialogueData,
    DialogueMetadata,
    DialogueStage,
    DialogueMessage,
)


def parse_teacher_doc(path: str) -> TeacherDocument:
    """
    解析教师文档

    支持格式：.docx, .md, .txt

    Args:
        path: 文件路径

    Returns:
        TeacherDocument: 解析后的教师文档
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    # 根据扩展名选择解析方式
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        if not HAS_DOCX:
            raise ImportError("请安装python-docx: pip install python-docx")
        raw_text = _parse_docx(file_path)
    elif suffix in [".md", ".txt", ".markdown"]:
        raw_text = file_path.read_text(encoding="utf-8")
    else:
        # 尝试作为文本文件读取
        raw_text = file_path.read_text(encoding="utf-8")

    # 提取结构化信息
    return TeacherDocument(
        raw_text=raw_text,
        teaching_objectives=_extract_objectives(raw_text),
        key_points=_extract_key_points(raw_text),
        workflow=_extract_workflow(raw_text),
        scoring_standard=_extract_scoring_standard(raw_text),
    )


def _parse_docx(file_path: Path) -> str:
    """解析docx文件"""
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def _extract_objectives(text: str) -> list[str]:
    """提取教学目标"""
    objectives = []

    # 匹配"教学目标"章节后的内容
    patterns = [
        r"教学目标[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|知识点|$)",
        r"教学目的[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|知识点|$)",
        r"目标[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|知识点|$)",
        r"学生能够[：:]?\s*\n?([^#]+?)(?=\n\s*[#\d]|$)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            content = match.group(1)
            # 按行分割并清理
            for line in content.split("\n"):
                line = line.strip()
                # 移除序号前缀
                line = re.sub(r"^\d+[\.、)）]\s*", "", line)
                line = re.sub(r"^[-•*]\s*", "", line)
                if line and len(line) > 3:
                    objectives.append(line)
            break

    return objectives


def _extract_key_points(text: str) -> list[str]:
    """提取知识点"""
    key_points = []

    patterns = [
        r"知识点[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|评分标准|$)",
        r"知识要点[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|评分标准|$)",
        r"核心概念[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|评分标准|$)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            content = match.group(1)
            for line in content.split("\n"):
                line = line.strip()
                line = re.sub(r"^\d+[\.、)）]\s*", "", line)
                line = re.sub(r"^[-•*]\s*", "", line)
                if line and len(line) > 3:
                    key_points.append(line)
            break

    return key_points


def _extract_workflow(text: str) -> Optional[list[str]]:
    """提取教学流程"""
    workflow = []

    patterns = [
        r"教学流程[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|知识点|$)",
        r"教学环节[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|知识点|$)",
        r"流程[：:]\s*\n?([^#]+?)(?=\n\s*[#\d]|教学目标|知识点|$)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            content = match.group(1)
            for line in content.split("\n"):
                line = line.strip()
                line = re.sub(r"^\d+[\.、)）]\s*", "", line)
                line = re.sub(r"^[-•*]\s*", "", line)
                if line and len(line) > 2:
                    workflow.append(line)
            break

    return workflow if workflow else None


def _extract_scoring_standard(text: str) -> Optional[str]:
    """提取评分标准"""
    patterns = [
        r"评分标准[：:]\s*\n?(.+?)(?=\n\s*[#]|$)",
        r"评价标准[：:]\s*\n?(.+?)(?=\n\s*[#]|$)",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(1).strip()

    return None


def parse_dialogue(path: str) -> DialogueData:
    """
    解析对话日志

    支持格式：.json, .txt

    Args:
        path: 文件路径

    Returns:
        DialogueData: 解析后的对话数据
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    suffix = file_path.suffix.lower()

    if suffix == ".json":
        return _parse_dialogue_json(file_path)
    else:
        # 默认按txt格式解析
        return _parse_dialogue_txt(file_path)


def _parse_dialogue_json(file_path: Path) -> DialogueData:
    """解析JSON格式对话日志"""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 提取元数据
    metadata = DialogueMetadata(
        task_id=data.get("task_id") or data.get("taskId"),
        profile=data.get("profile") or data.get("student_profile"),
        start_time=data.get("start_time"),
        end_time=data.get("end_time"),
        total_rounds=data.get("total_rounds") or data.get("round_count"),
    )

    # 提取对话阶段
    stages = []
    stages_data = data.get("stages", [])

    if stages_data:
        for stage_data in stages_data:
            messages = []
            for msg in stage_data.get("messages", []):
                messages.append(DialogueMessage(
                    role=msg.get("role", "unknown"),
                    content=msg.get("content", ""),
                    timestamp=msg.get("timestamp"),
                ))

            stages.append(DialogueStage(
                stage_id=str(stage_data.get("stage_id", "")),
                stage_name=stage_data.get("stage_name", ""),
                messages=messages,
            ))
    else:
        # 尝试从messages字段解析
        messages_data = data.get("messages", [])
        if messages_data:
            messages = []
            for msg in messages_data:
                messages.append(DialogueMessage(
                    role=msg.get("role", "unknown"),
                    content=msg.get("content", ""),
                    timestamp=msg.get("timestamp"),
                ))
            stages.append(DialogueStage(
                stage_id="default",
                stage_name="对话",
                messages=messages,
            ))

    return DialogueData(
        metadata=metadata,
        stages=stages,
        raw_text=json.dumps(data, ensure_ascii=False, indent=2),
    )


def _parse_dialogue_txt(file_path: Path) -> DialogueData:
    """解析TXT格式对话日志"""
    text = file_path.read_text(encoding="utf-8")

    messages = []
    lines = text.split("\n")

    # 尝试提取元数据
    task_id = None
    profile = None

    for line in lines[:20]:  # 检查前20行
        if "任务ID" in line or "task_id" in line.lower():
            match = re.search(r"[:：]\s*(\S+)", line)
            if match:
                task_id = match.group(1)
        if "学生画像" in line or "profile" in line.lower():
            match = re.search(r"[:：]\s*(\S+)", line)
            if match:
                profile = match.group(1)

    # 解析对话消息
    # 支持格式：
    # [学生] xxx
    # AI: xxx
    # 学生: xxx
    # === 第N轮 ===

    current_role = None
    current_content = []

    role_patterns = [
        (r"^\[?学生\]?[：:]\s*", "student"),
        (r"^\[?AI\]?[：:]\s*", "ai"),
        (r"^\[?老师\]?[：:]\s*", "ai"),
        (r"^\[?系统\]?[：:]\s*", "system"),
    ]

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 检查是否是新的角色行
        matched = False
        for pattern, role in role_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                # 保存之前的消息
                if current_role and current_content:
                    messages.append(DialogueMessage(
                        role=current_role,
                        content="\n".join(current_content).strip(),
                    ))

                # 开始新消息
                current_role = role
                current_content = [re.sub(pattern, "", line, flags=re.IGNORECASE)]
                matched = True
                break

        if not matched:
            # 可能是轮次分隔符
            if re.match(r"^[=\-#]+.*[=\-#]+$", line):
                if current_role and current_content:
                    messages.append(DialogueMessage(
                        role=current_role,
                        content="\n".join(current_content).strip(),
                    ))
                    current_role = None
                    current_content = []
            elif current_role:
                current_content.append(line)

    # 保存最后一条消息
    if current_role and current_content:
        messages.append(DialogueMessage(
            role=current_role,
            content="\n".join(current_content).strip(),
        ))

    # 统计轮数
    student_count = sum(1 for m in messages if m.role == "student")

    metadata = DialogueMetadata(
        task_id=task_id,
        profile=profile,
        total_rounds=student_count,
    )

    return DialogueData(
        metadata=metadata,
        stages=[DialogueStage(
            stage_id="default",
            stage_name="对话",
            messages=messages,
        )],
        raw_text=text,
    )
